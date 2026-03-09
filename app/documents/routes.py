# app/documents/routes.py
from __future__ import annotations

import json
import mimetypes
import shutil
import subprocess
import tempfile
import time as time_mod
from datetime import date, datetime, time, timedelta
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from docx.shared import Pt
from flask import abort, current_app, flash, redirect, render_template, request, send_file, url_for
from flask_login import current_user, login_required
from werkzeug.utils import secure_filename

from app.extensions import db
from app.models import Category, Document, GoogleToken
from app.utils.dates import (
    ddmmyyyy_to_slash as _ddmmyyyy_to_slash,
    parse_date_ddmmyyyy as _parse_due_date_date,
    parse_datetime_ddmmyyyy as _parse_date_ddmmyyyy,
)
from . import bp
from .forms import DocumentEditForm, FillOficioForm


# =========================================================
# Config de plantillas de Oficios (multi-plantilla)
# =========================================================
OFICIO_TEMPLATES: dict[str, dict] = {
    "oficio_respuesta": {
        "display_name": "Oficio Respuesta",
        "filename": "oficio_respuesta_template_v1.docx",
        "fields": [
            "numero_solicitud",
            "fecha_solicitud",
            "de_nombre",
            "de_cargo",
            "a_nombre",
            "tenor_literal",
            "respuesta",
        ],
        "mapping": {
            "{{NUM_SOLICITUD}}": "numero_solicitud",
            "{{FECHA_SOLICITUD}}": "fecha_solicitud",
            "{{DE_NOMBRE}}": "de_nombre",
            "{{DE_CARGO}}": "de_cargo",
            "{{A_NOMBRE}}": "a_nombre",
            "{{TENOR_LITERAL}}": "tenor_literal",
            "{{RESPUESTA}}": "respuesta",
        },
    },
}


# =========================================================
# Helpers
# =========================================================
def _allowed_file(filename: str) -> bool:
    allowed = current_app.config.get("ALLOWED_EXTENSIONS") or set()
    if not allowed:
        return True
    if "." not in filename:
        return False
    ext = filename.rsplit(".", 1)[1].lower()
    return ext in allowed


def _google_connected(user_id: int) -> bool:
    try:
        return bool(GoogleToken.query.filter_by(user_id=user_id).first())
    except Exception:
        return False


def _create_calendar_deadline_best_effort(title: str, due: date) -> bool:
    """
    Best-effort: crea evento a las 09:00 hora local (1h duración).
    - No rompe el flujo si falla.
    - Soporta create_deadline_event que retorne:
      - bool
      - (bool, str)
    """
    try:
        from app.calendar_bp.google_service import create_deadline_event  # type: ignore
    except Exception:
        return False

    start_dt = datetime.combine(due, time(9, 0))
    end_dt = start_dt + timedelta(hours=1)

    description = (
        "Recordatorio creado desde Repositorio Municunco.\n"
        "Tip: Mantén los documentos al día para auditoría y trazabilidad."
    )

    try:
        res = create_deadline_event(
            user_id=current_user.id,
            title=(title or "")[:120],
            description=description,
            start_dt=start_dt,
            end_dt=end_dt,
        )

        if isinstance(res, tuple) and len(res) >= 1:
            ok = bool(res[0])
            err = ""
            if len(res) >= 2 and isinstance(res[1], str):
                err = res[1]
            if not ok and err:
                current_app.logger.warning("Calendar reminder failed: %s", err)
            return ok

        return bool(res)

    except Exception as e:
        current_app.logger.warning("Calendar best-effort exception: %s", e)
        return False


def _is_previewable(filename: str) -> bool:
    fn = (filename or "").lower()
    return fn.endswith((".pdf", ".png", ".jpg", ".jpeg", ".webp"))


def _safe_slug(value: str) -> str:
    raw = (value or "").strip()
    out = []
    for ch in raw:
        if ch.isalnum() or ch in ("-", "_"):
            out.append(ch)
    return "".join(out)[:60] or "generado"


def _generated_dir() -> Path:
    base = Path(current_app.config["UPLOAD_PATH"])
    gen = base / "generated"
    gen.mkdir(parents=True, exist_ok=True)
    return gen


def _templates_dir() -> Path:
    base = Path(current_app.config["UPLOAD_PATH"])
    configured = (current_app.config.get("DOCX_TEMPLATES_DIR") or "").strip()
    if configured:
        return Path(configured)
    return base / "templates"


def _oficio_template_path(filename: str) -> Path:
    return _templates_dir() / filename


def _ensure_default_font(doc: DocxDocument, font_name: str = "Arial", font_size_pt: int = 11) -> None:
    try:
        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size_pt)
    except Exception:
        pass


def _replace_in_paragraph_runs(paragraph, mapping: dict[str, str]) -> None:
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    changed = False
    for k, v in mapping.items():
        if k in full_text:
            full_text = full_text.replace(k, v)
            changed = True

    if not changed:
        return

    paragraph.runs[0].text = full_text
    for r in paragraph.runs[1:]:
        r.text = ""


def _docx_replace_text(doc: DocxDocument, mapping: dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph_runs(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph_runs(p, mapping)


def _docx_force_bold_placeholders(
    doc: DocxDocument,
    keys: set[str],
    font_name: str = "Arial",
    font_size_pt: int = 11,
) -> None:
    def _apply_on_paragraph(paragraph) -> None:
        for run in paragraph.runs:
            if not run.text:
                continue
            for k in keys:
                if k in run.text:
                    try:
                        run.bold = True
                        run.font.name = font_name
                        run.font.size = Pt(font_size_pt)
                    except Exception:
                        pass
                    break

    for p in doc.paragraphs:
        _apply_on_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _apply_on_paragraph(p)


def _convert_docx_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)

    lo_bin = (current_app.config.get("LIBREOFFICE_BIN") or "soffice").strip()
    lo_real = shutil.which(lo_bin) or shutil.which("soffice")
    if lo_real:
        cmd = [
            lo_real,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(docx_path),
        ]
        proc = subprocess.run(cmd, capture_output=True, text=True)
        if proc.returncode != 0:
            raise RuntimeError(f"LibreOffice falló al convertir a PDF: {proc.stderr or proc.stdout}")

        pdf_path = out_dir / f"{docx_path.stem}.pdf"
        if not pdf_path.exists():
            raise RuntimeError("No se generó el PDF (salida inesperada de LibreOffice).")
        return pdf_path

    try:
        from docx2pdf import convert  # type: ignore
    except Exception:
        raise RuntimeError(
            "No hay convertidor disponible. En servidor Ubuntu instala LibreOffice. "
            "En Windows instala docx2pdf + Microsoft Word."
        )

    pdf_path = out_dir / f"{docx_path.stem}.pdf"

    try:
        import pythoncom  # type: ignore

        pythoncom.CoInitialize()
        coinit = True
    except Exception:
        coinit = False

    try:
        convert(str(docx_path), str(pdf_path))
    except Exception as e:
        raise RuntimeError(f"docx2pdf falló (¿Word instalado?): {e}")
    finally:
        if coinit:
            try:
                import pythoncom  # type: ignore

                pythoncom.CoUninitialize()
            except Exception:
                pass

    for _ in range(30):
        if pdf_path.exists():
            try:
                with open(pdf_path, "rb"):
                    pass
                return pdf_path
            except OSError:
                time_mod.sleep(0.15)

    if not pdf_path.exists():
        raise RuntimeError("No se generó el PDF con docx2pdf.")

    raise RuntimeError("El PDF fue generado pero quedó bloqueado por Word. Reintenta en unos segundos.")


def _send_pdf_from_temp(pdf_path: Path, download_name: str):
    tmp = tempfile.NamedTemporaryFile(prefix="municunco_dl_", suffix=".pdf", delete=False)
    tmp_path = Path(tmp.name)
    tmp.close()

    last_err: Exception | None = None
    for _ in range(20):
        try:
            shutil.copyfile(pdf_path, tmp_path)
            last_err = None
            break
        except Exception as e:
            last_err = e
            time_mod.sleep(0.15)

    if last_err is not None:
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except Exception:
            pass
        raise last_err

    resp = send_file(
        tmp_path,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=download_name,
    )

    @resp.call_on_close
    def _cleanup():
        try:
            tmp_path.unlink(missing_ok=True)
        except TypeError:
            try:
                if tmp_path.exists():
                    tmp_path.unlink()
            except Exception:
                pass
        except Exception:
            pass

    return resp


def _build_oficio_defs_for_front() -> dict[str, dict]:
    """
    Solo lo que el front necesita:
    - display_name
    - fields
    - display_path
    """
    out: dict[str, dict] = {}
    for k, v in OFICIO_TEMPLATES.items():
        fn = v.get("filename") or ""
        out[k] = {
            "display_name": v.get("display_name") or k,
            "fields": v.get("fields") or [],
            "display_path": f"uploads/templates/{fn}" if fn else "—",
        }
    return out


def _setup_oficio_form_choices(form: FillOficioForm) -> None:
    form.oficio_tipo.choices = [(k, v["display_name"]) for k, v in OFICIO_TEMPLATES.items()]

    cats = Category.query.order_by(Category.name.asc()).all()
    form.category_id.choices = [(0, "— Selecciona una categoría —")] + [(c.id, c.name) for c in cats]


def _setup_document_edit_form_choices(form: DocumentEditForm) -> None:
    cats = Category.query.order_by(Category.name.asc()).all()
    form.category_id.choices = [(0, "— Sin categoría —")] + [(c.id, c.name) for c in cats]


def _delete_document_file(doc: Document) -> None:
    path = Path(doc.path or "")
    try:
        if path.exists() and path.is_file():
            path.unlink()
    except Exception as ex:
        current_app.logger.warning("No se pudo eliminar archivo físico del documento %s: %s", doc.id, ex)


def _can_manage_documents() -> bool:
    return bool(current_user.is_admin)


# =========================================================
# Rutas
# =========================================================
@bp.get("/dashboard")
@login_required
def dashboard():
    total_docs = Document.query.count()
    total_cats = Category.query.count()

    used_bytes = db.session.query(db.func.coalesce(db.func.sum(Document.file_size), 0)).scalar() or 0
    used_mb = round((used_bytes / 1024 / 1024), 2)

    last_doc = Document.query.order_by(Document.created_at.desc()).first()
    last_doc_name = (last_doc.filename or last_doc.name) if last_doc else "Ninguno"

    gc_connected = _google_connected(current_user.id)

    today = date.today()
    until = today + timedelta(days=7)

    due_soon_docs = (
        Document.query.filter(Document.due_date.isnot(None))
        .filter(Document.due_date >= today)
        .filter(Document.due_date <= until)
        .order_by(Document.due_date.asc())
        .limit(8)
        .all()
    )

    overdue_count = (
        Document.query.filter(Document.due_date.isnot(None))
        .filter(Document.due_date < today)
        .count()
    )

    deadlines_7d: list[dict[str, str | None]] = []

    if gc_connected:
        try:
            from app.calendar_bp.google_service import list_events_range  # type: ignore

            start_dt = datetime.combine(today, time.min)
            end_dt = datetime.combine(until, time.max)

            raw_events = list_events_range(
                user_id=current_user.id,
                start_dt=start_dt,
                end_dt=end_dt,
                max_results=50,
            )

            deadlines_7d = [
                {
                    "id": ev.get("id"),
                    "summary": ev.get("summary") or "(sin título)",
                    "when": ev.get("when") or "—",
                    "link": ev.get("link"),
                }
                for ev in raw_events[:8]
            ]

        except Exception as e:
            current_app.logger.warning("Dashboard calendar fetch failed: %s", e)
            deadlines_7d = []

    return render_template(
        "dashboard.html",
        total_docs=total_docs,
        total_cats=total_cats,
        used_mb=used_mb,
        last_doc_name=last_doc_name or "Ninguno",
        gc_connected=gc_connected,
        due_soon_docs=due_soon_docs,
        overdue_count=overdue_count,
        deadlines_7d=deadlines_7d,
    )


@bp.get("/")
@login_required
def index():
    q = (request.args.get("q") or "").strip()
    cat_id = (request.args.get("category_id") or "").strip()

    desde_raw = request.args.get("desde") or ""
    hasta_raw = request.args.get("hasta") or ""
    desde_dt = _parse_date_ddmmyyyy(desde_raw)
    hasta_dt = _parse_date_ddmmyyyy(hasta_raw)

    due_status = (request.args.get("due_status") or "all").strip()
    sort = (request.args.get("sort") or "created_desc").strip()

    query = Document.query

    if q:
        like = f"%{q.lower()}%"
        query = query.filter(db.func.lower(Document.name).like(like))

    if cat_id.isdigit():
        query = query.filter(Document.category_id == int(cat_id))

    if desde_dt:
        query = query.filter(Document.created_at >= datetime.combine(desde_dt.date(), time.min))
    if hasta_dt:
        query = query.filter(Document.created_at <= datetime.combine(hasta_dt.date(), time.max))

    today = date.today()
    if due_status == "has_due":
        query = query.filter(Document.due_date.isnot(None))
    elif due_status == "no_due":
        query = query.filter(Document.due_date.is_(None))
    elif due_status == "overdue":
        query = query.filter(Document.due_date.isnot(None)).filter(Document.due_date < today)
    elif due_status == "due_soon":
        query = (
            query.filter(Document.due_date.isnot(None))
            .filter(Document.due_date >= today)
            .filter(Document.due_date <= (today + timedelta(days=7)))
        )

    if sort == "due_asc":
        query = query.order_by(
            db.case((Document.due_date.is_(None), 1), else_=0),
            Document.due_date.asc(),
            Document.created_at.desc(),
        )
    elif sort == "due_desc":
        query = query.order_by(
            db.case((Document.due_date.is_(None), 1), else_=0),
            Document.due_date.desc(),
            Document.created_at.desc(),
        )
    else:
        query = query.order_by(Document.created_at.desc())

    docs = query.all()
    categories = Category.query.order_by(Category.name.asc()).all()

    return render_template(
        "documents/index.html",
        docs=docs,
        categories=categories,
        q=q,
        category_id=cat_id,
        desde=desde_raw,
        hasta=hasta_raw,
        due_status=due_status,
        sort=sort,
        can_manage_docs=_can_manage_documents(),
    )


@bp.get("/upload")
@login_required
def upload():
    categories = Category.query.order_by(Category.name.asc()).all()
    gc_connected = _google_connected(current_user.id)
    return render_template("documents/upload.html", categories=categories, gc_connected=gc_connected)


@bp.post("/upload")
@login_required
def upload_post():
    files = request.files.getlist("files")
    category_id_raw = (request.form.get("category_id") or "").strip()
    category_id = int(category_id_raw) if category_id_raw.isdigit() else None

    due_raw = request.form.get("due_date") or ""
    due = _parse_due_date_date(due_raw)

    if due and due < date.today():
        flash("La fecha límite no puede ser en el pasado.", "warning")
        return redirect(url_for("documents.upload"))

    if not files or all((f is None or not f.filename) for f in files):
        flash("Selecciona al menos un archivo.", "warning")
        return redirect(url_for("documents.upload"))

    upload_path = Path(current_app.config["UPLOAD_PATH"])
    upload_path.mkdir(parents=True, exist_ok=True)

    saved = 0
    rejected = 0
    created_calendar = 0
    calendar_failed = 0

    gc_connected = _google_connected(current_user.id)

    for f in files:
        if not f or not f.filename:
            continue

        original = secure_filename(f.filename)
        if not original or not _allowed_file(original):
            rejected += 1
            continue

        ext = "." + original.rsplit(".", 1)[1].lower() if "." in original else ""
        storage_name = f"{uuid4().hex}{ext}"
        disk_path = upload_path / storage_name

        f.save(disk_path)
        size = disk_path.stat().st_size if disk_path.exists() else 0

        doc = Document(
            name=Path(original).stem[:160] or original[:160],
            filename=original[:260],
            path=str(disk_path),
            file_size=int(size),
            category_id=category_id,
            uploaded_by_id=current_user.id,
            due_date=due,
        )
        db.session.add(doc)
        saved += 1

        if due and gc_connected:
            ok = _create_calendar_deadline_best_effort(title=doc.filename, due=due)
            if ok:
                created_calendar += 1
            else:
                calendar_failed += 1

    db.session.commit()

    if saved:
        flash(f"✅ Subida completada: {saved} archivo(s).", "success")
    if rejected:
        flash(f"⚠️ {rejected} archivo(s) rechazado(s) por nombre/extensión.", "warning")

    if due and gc_connected:
        if created_calendar:
            flash(f"🗓️ Recordatorios creados en Google Calendar: {created_calendar}.", "info")
        if calendar_failed:
            flash(f"⚠️ No se pudieron crear {calendar_failed} recordatorio(s) en Google Calendar.", "warning")

    return redirect(url_for("documents.index"))


@bp.get("/edit/<int:doc_id>")
@login_required
def edit(doc_id: int):
    if not _can_manage_documents():
        abort(403)

    doc = Document.query.get_or_404(doc_id)

    form = DocumentEditForm()
    _setup_document_edit_form_choices(form)

    form.name.data = doc.name or ""
    form.category_id.data = doc.category_id or 0
    form.due_date.data = doc.due_date.strftime("%d-%m-%Y") if doc.due_date else ""
    form.note.data = ""

    return render_template("documents/edit.html", form=form, doc=doc)


@bp.post("/edit/<int:doc_id>")
@login_required
def edit_post(doc_id: int):
    if not _can_manage_documents():
        abort(403)

    doc = Document.query.get_or_404(doc_id)

    form = DocumentEditForm()
    _setup_document_edit_form_choices(form)

    if not form.validate_on_submit():
        return render_template("documents/edit.html", form=form, doc=doc), 400

    due_raw = (form.due_date.data or "").strip()
    due = _parse_due_date_date(due_raw)

    if due_raw and not due:
        form.due_date.errors.append("Ingresa una fecha válida en formato dd-mm-aaaa.")
        return render_template("documents/edit.html", form=form, doc=doc), 400

    doc.name = (form.name.data or "").strip()[:160]
    doc.category_id = int(form.category_id.data or 0) or None
    doc.due_date = due

    db.session.commit()

    flash("Documento actualizado correctamente.", "success")
    return redirect(url_for("documents.index"))


@bp.get("/preview/<int:doc_id>")
@login_required
def preview(doc_id: int):
    doc = Document.query.get_or_404(doc_id)
    path = Path(doc.path)
    if not path.exists():
        abort(404, description="Archivo no encontrado en disco.")

    if not _is_previewable(doc.filename):
        abort(415, description="Vista previa disponible solo para PDF e imágenes.")

    mime, _ = mimetypes.guess_type(doc.filename)
    return send_file(path, mimetype=mime or "application/octet-stream", as_attachment=False, download_name=doc.filename)


@bp.get("/download/<int:doc_id>")
@login_required
def download(doc_id: int):
    doc = Document.query.get_or_404(doc_id)
    path = Path(doc.path)
    if not path.exists():
        flash("Archivo no encontrado en disco.", "danger")
        return redirect(url_for("documents.index"))

    mime, _ = mimetypes.guess_type(doc.filename)
    return send_file(path, mimetype=mime or "application/octet-stream", as_attachment=True, download_name=doc.filename)


@bp.post("/delete/<int:doc_id>")
@login_required
def delete(doc_id: int):
    if not _can_manage_documents():
        abort(403)

    doc = Document.query.get_or_404(doc_id)

    db.session.delete(doc)
    db.session.commit()

    _delete_document_file(doc)

    flash("Documento eliminado.", "success")
    return redirect(url_for("documents.index"))


@bp.post("/delete-many")
@login_required
def delete_many():
    if not _can_manage_documents():
        abort(403)

    raw_ids = request.form.getlist("doc_ids")
    doc_ids: list[int] = []

    for raw in raw_ids:
        raw = (raw or "").strip()
        if raw.isdigit():
            value = int(raw)
            if value > 0:
                doc_ids.append(value)

    doc_ids = list(dict.fromkeys(doc_ids))

    if not doc_ids:
        flash("Selecciona al menos un documento para eliminar.", "warning")
        return redirect(url_for("documents.index"))

    docs = Document.query.filter(Document.id.in_(doc_ids)).all()
    if not docs:
        flash("No se encontraron documentos válidos para eliminar.", "warning")
        return redirect(url_for("documents.index"))

    deleted_count = 0
    deleted_files: list[Document] = []

    for doc in docs:
        deleted_files.append(doc)
        db.session.delete(doc)
        deleted_count += 1

    db.session.commit()

    for doc in deleted_files:
        _delete_document_file(doc)

    flash(f"Se eliminaron {deleted_count} documento(s) correctamente.", "success")
    return redirect(url_for("documents.index"))


# =========================================================
# Oficios (multi-plantilla)
# URL: /documents/oficio
# =========================================================
@bp.get("/oficio")
@login_required
def oficio():
    form = FillOficioForm()
    _setup_oficio_form_choices(form)

    default_key = next(iter(OFICIO_TEMPLATES.keys()))
    form.oficio_tipo.data = default_key

    form.numero_solicitud.data = "MU071T0001762"
    form.fecha_solicitud.data = "20-01-2026"
    form.de_nombre.data = "NELSON OLIVERA STAUB"
    form.de_cargo.data = "ADMINISTRADOR MUNICIPAL"
    form.a_nombre.data = "CRISTINA INOSTROZA DELGADO"
    form.tenor_literal.data = "texto de prueba"
    form.respuesta.data = "respuesta de prueba"
    form.guardar_pdf.data = True
    form.category_id.data = 0

    defs = _build_oficio_defs_for_front()
    selected = defs.get(default_key, {})
    return render_template(
        "documents/oficio.html",
        form=form,
        oficio_defs_json=json.dumps(defs, ensure_ascii=False),
        selected_template_path=selected.get("display_path", "—"),
    )


@bp.post("/oficio")
@login_required
def oficio_post():
    form = FillOficioForm()
    _setup_oficio_form_choices(form)

    if not form.validate_on_submit():
        defs = _build_oficio_defs_for_front()
        selected_key = form.oficio_tipo.data or next(iter(OFICIO_TEMPLATES.keys()))
        selected = defs.get(selected_key, {})
        return (
            render_template(
                "documents/oficio.html",
                form=form,
                oficio_defs_json=json.dumps(defs, ensure_ascii=False),
                selected_template_path=selected.get("display_path", "—"),
            ),
            400,
        )

    tipo = (form.oficio_tipo.data or "").strip()
    if tipo not in OFICIO_TEMPLATES:
        flash("Tipo de oficio inválido.", "danger")
        return redirect(url_for("documents.oficio"))

    tpl_def = OFICIO_TEMPLATES[tipo]
    tpl_filename = tpl_def.get("filename") or ""
    tpl_path = _oficio_template_path(tpl_filename)
    if not tpl_path.exists():
        flash(f"No se encontró la plantilla seleccionada en: {tpl_path}", "danger")
        return redirect(url_for("documents.oficio"))

    mapping: dict[str, str] = {}

    for placeholder, field_name in (tpl_def.get("mapping") or {}).items():
        val = getattr(form, field_name).data if hasattr(form, field_name) else ""
        val = (val or "").strip()

        if field_name == "fecha_solicitud":
            val = _ddmmyyyy_to_slash(val)

        mapping[str(placeholder)] = val

    numero = (form.numero_solicitud.data or "").strip()
    safe_num = _safe_slug(numero)

    guardar_pdf = bool(form.guardar_pdf.data)
    category_id = int(form.category_id.data or 0) if guardar_pdf else 0
    final_category_id = category_id if (guardar_pdf and category_id > 0) else None

    try:
        with tempfile.TemporaryDirectory(prefix="municunco_oficio_") as tmp:
            tmp_dir = Path(tmp)

            d = DocxDocument(str(tpl_path))
            _ensure_default_font(d, "Arial", 11)

            _docx_replace_text(d, mapping)

            _docx_force_bold_placeholders(
                d,
                keys={"{{DE_NOMBRE}}", "{{DE_CARGO}}", "{{A_NOMBRE}}"},
                font_name="Arial",
                font_size_pt=11,
            )

            out_docx = tmp_dir / f"oficio_{safe_num}.docx"
            d.save(str(out_docx))

            tmp_pdf = _convert_docx_to_pdf(out_docx, tmp_dir)

            if guardar_pdf:
                gen_dir = _generated_dir()
                final_name = f"Oficio_{safe_num}_{uuid4().hex[:8]}.pdf"
                final_path = gen_dir / final_name
                shutil.copyfile(tmp_pdf, final_path)

                size = final_path.stat().st_size if final_path.exists() else 0
                new_doc = Document(
                    name=f"Oficio_{safe_num}"[:160],
                    filename=final_name[:260],
                    path=str(final_path),
                    file_size=int(size),
                    category_id=final_category_id,
                    uploaded_by_id=current_user.id,
                    due_date=None,
                )
                db.session.add(new_doc)
                db.session.commit()

                flash("✅ PDF generado y guardado en el repositorio.", "success")

                return send_file(
                    final_path,
                    mimetype="application/pdf",
                    as_attachment=True,
                    download_name=final_name,
                )

            return _send_pdf_from_temp(
                tmp_pdf,
                download_name=f"Oficio_{safe_num}.pdf",
            )

    except Exception as e:
        current_app.logger.exception("Oficio generation failed: %s", e)
        flash(f"No se pudo generar el PDF: {e}", "danger")

        defs = _build_oficio_defs_for_front()
        selected = defs.get(tipo, {})
        return (
            render_template(
                "documents/oficio.html",
                form=form,
                oficio_defs_json=json.dumps(defs, ensure_ascii=False),
                selected_template_path=selected.get("display_path", "—"),
            ),
            500,
        )